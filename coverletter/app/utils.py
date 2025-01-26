from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

client = OpenAI(api_key=api_key)

def extract_text_from_file(file_path):
    """Extract text from PDF or DOCX files."""
    if file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        text = " ".join(page.extract_text() for page in reader.pages if page.extract_text())
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        text = " ".join(para.text for para in doc.paragraphs)
    else:
        raise ValueError("Unsupported file format.")
    return text

def generate_cover_letter(job_description, cv_text):
    messages = [
        {"role": "system", "content":  (
            "You are a helpful assistant that writes professional cover letters."
            "Keep in mind the following guidelines:\n\n"
            "1. Use a gender-neutral opening salutation, such as 'Dear Hiring Manager' or 'Dear Recruiting Team'.\n"
            "2. Ensure all necessary contact details are included.\n"
            "3. Limit the body to 250-450 words.\n"
            "4. Structure the letter into 3-4 paragraphs: Introduction, Body, and Closing.\n"
            "5. Demonstrate alignment of personal values and goals with the company's mission.\n"
            "6. Use numbers and metrics to quantify achievements, applying the STAR method (Situation, Task, Action, Result).\n"
            "7. Avoid vague language like 'as well as' or 'and more'; instead, use strong, action-oriented words.\n"
            "8. Maintain a formal and professional tone throughout.\n"
            "9. Conclude with a professional closing salutation. For example, 'Sincerely, [Full name]'\n"
            "10 Include contact details found on the CV.\n"
            "11 Do not change any of the contact details found on the CV and keep them in their original spelling"
        )},
        {"role": "user", "content": f"Write a professional cover letter tailored to the following job description:\n\n{job_description}\n\nBased on the following CV:\n\n{cv_text}"}
    ]

    response = client.chat.completions.create(
        model="gpt-4",
        messages=messages,
        max_tokens=500,
    )

    # Extract the generated text from the response
    cover_letter = response.choices[0].message.content
    cover_letter = cover_letter.strip()
    
    return cover_letter

def create_pdf(text, file_path):
    """Generate a well-formatted PDF with the provided text."""
    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter

    # Configure font and line spacing
    c.setFont("Helvetica", 12)
    margin = 40
    line_height = 14
    max_width = width - 2 * margin
    y_position = height - margin  # Start near the top of the page

    # Split text into paragraphs
    paragraphs = text.split("\n")

    for paragraph in paragraphs:
        # Wrap text to fit within the page margins
        lines = simpleSplit(paragraph, "Helvetica", 12, max_width)
        for line in lines:
            if y_position < margin + line_height:  # Check if we need a new page
                c.showPage()
                c.setFont("Helvetica", 12)
                y_position = height - margin
            c.drawString(margin, y_position, line)
            y_position -= line_height  # Move down for the next line

        # Add spacing between paragraphs
        y_position -= line_height

    c.save()


def create_docx(text, file_path):
    """Generate a DOCX with the provided text."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)
